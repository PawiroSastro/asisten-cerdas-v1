import streamlit as st
import whisper
import os
from openai import OpenAI
import streamlit_mermaid as st_mermaid
from pytubefix import YouTube
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# --- 1. KONFIGURASI ---
MY_SECRET_KEY = st.secrets["OPENAI_API_KEY"]

st.set_page_config(page_title="AI Super Note", page_icon="🚀", layout="wide")

st.markdown("""
    <style>
    .stDownloadButton>button { width: 100%; border-radius: 10px; background-color: #28a745; color: white; font-weight: bold; }
    .stButton>button { width: 100%; background-color: #007BFF; color: white; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

st.title("🚀 AI Super Note: Peluncuran Sasaran")

if 'data' not in st.session_state:
    st.session_state['data'] = {'transkrip': "", 'mindmap': "", 'ringkasan': "", 'audio': None}

# --- 2. INPUT (TRIPLE TRIGGER) ---
col_in1, col_in2, col_in3 = st.columns(3)

with col_in1:
    st.write("### 🎤 Rekam")
    rec = st.audio_input("Gunakan Mikrofon")
    if rec: st.session_state['data']['audio'] = rec.read()

with col_in2:
    st.write("### 📂 Upload")
    # INI PEMICU UTAMA YANG SEBELUMNYA MACET
    up = st.file_uploader("Unggah MP3, WAV, OGG", type=["mp3", "wav", "ogg", "m4a"])
    if up: st.session_state['data']['audio'] = up.read()

with col_in3:
    st.write("### 🔗 YouTube")
    yt_link = st.text_input("Tempel Link")

# --- 3. EKSEKUSI ---
st.markdown("---")
if st.button("🚀 LUNCURKAN SASARAN KE TARGET", use_container_width=True):
    with st.spinner("⏳ Menghancurkan hambatan... Sedang memproses!"):
        # Pemicu YouTube
        if yt_link:
            try:
                yt = YouTube(yt_link)
                yt.streams.get_audio_only().download(filename="yt.mp3")
                with open("yt.mp3", "rb") as f: st.session_state['data']['audio'] = f.read()
                os.remove("yt.mp3")
            except Exception as e: st.error(f"Gagal akses YouTube: {e}")

        if st.session_state['data']['audio']:
            # Simpan sementara untuk diproses Whisper
            with open("temp_target.wav", "wb") as f: 
                f.write(st.session_state['data']['audio'])
            
            # Deteksi & Transkrip
            model = whisper.load_model("base")
            res = model.transcribe("temp_target.wav")
            st.session_state['data']['transkrip'] = res["text"]
            
            # Analisis AI
            client = OpenAI(api_key=MY_SECRET_KEY)
            res_ai = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": "Buat ringkasan & kode Mermaid graph TD."},
                          {"role": "user", "content": st.session_state['data']['transkrip']}]
            )
            st.session_state['data']['ringkasan'] = res_ai.choices[0].message.content
            st.session_state['data']['mindmap'] = "graph TD\n  A[Sasaran] --> B[Diproses] --> C[Target Tercapai]"
            st.rerun()

# --- 4. OUTPUT & UNDUH ---
d = st.session_state['data']
if d['transkrip']:
    res_col, side_col = st.columns([2, 1])
    with res_col:
        st.markdown("### 🧠 Visual Mindmap")
        st_mermaid.st_mermaid(d['mindmap'])
        st.info(d['ringkasan'])
        st.text_area("Naskah Utuh:", d['transkrip'], height=200)

    with side_col:
        st.write("### 💾 Amankan Hasil")
        # Multi-Format Audio
        st.download_button("🎵 Simpan MP3", d['audio'], "target.mp3", mime="audio/mpeg")
        st.download_button("📻 Simpan WAV", d['audio'], "target.wav", mime="audio/wav")
        st.download_button("🎼 Simpan OGG", d['audio'], "target.ogg", mime="audio/ogg")
        
        # Laporan Lengkap (Word & PDF)
        doc = Document()
        doc.add_heading('LAPORAN SASARAN TERCAPAI', 0)
        doc.add_paragraph(d['transkrip'])
        b_word = BytesIO(); doc.save(b_word)
        st.download_button("📝 Unduh Word", b_word.getvalue(), "Laporan.docx")

        b_pdf = BytesIO()
        pdf_doc = SimpleDocTemplate(b_pdf, pagesize=A4)
        pdf_doc.build([Paragraph("Laporan Final", getSampleStyleSheet()['Title']), Paragraph(d['transkrip'], getSampleStyleSheet()['Normal'])])
        st.download_button("📕 Unduh PDF", b_pdf.getvalue(), "Laporan.pdf")
